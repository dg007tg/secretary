from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

FEE_NAMES = {
        "SALARY_FEE":"工资",
        "NURSING_FEE":"护理费"
    }


class GEN_REPORT:
    header_format = u"仁爱堂养老院%s年%s月财务报表"
    introduction = "  这是仁爱堂养老院的财务报表，报表将详细呈现本月的每一条消费记录\
以及这些记录汇总的总金额。"
    content_font_size = 14
    table_font_size = 12
    def __change_font(self,obj,font_name='宋体',size=None):
        obj.font.name = font_name
        obj._element.rPr.rFonts.set(qn('w:eastAsia'),font_name)
        if (size) and isinstance(size, Pt):
            obj.font.size = size

    def __init__(self,report):
        '''
        @param month the month to which the report correspond to
        '''
        self.doc = Document()
        self.month = self.__get_month(report)

        '''
        节是word格式化的最大单位，段落等概念都是节的一部分，文档默认有
        一个节。此处相当于设置整体文档的基础格式
        '''
        distance = Inches(1)
        sec = self.doc.sections[0]
        #设置页面边距
        sec.left_margin = distance
        sec.right_margin = distance
        sec.top_margin = distance
        sec.botton_margin = distance
        #设置页面宽高
        sec.page_width = Inches(12)
        sec.page_height = Inches(20)
        #设置默认字体
        self.__change_font(self.doc.styles['Normal'],font_name = "宋体")

        self.__add_header()
        self.__add_body(report)
        '''
        #add paragraph
        paragraph = self.doc.add_paragraph("仁爱堂养老院财务报表")
        para_format = paragraph.paragraph_format
        #设置段前距
        para_format.space_before = Pt(10)
        #段后距
        para_format.space_after = Pt(12)
        #行间距
        para_format.line_spacing = Pt(19)

        如果希望同一段落中的文本格式不同，就需要使用Run对象
        （可以理解为可以单独设置格式的段落内对象）。
        paragraph为一个doc.paragraph()对象
        run = paragraph.add_run("text...")
        run.bold = True
        ...
        change_font(run,...)

        '''

    def save(self,path):
        import re
        if re.search(r"\.docx$", path):
            self.doc.save(path)
            print("Report is save as %s successfully." % (path,))
        else:
            print("Illegal path to save the report!")

    def __get_month(self,report):
        '''
        Inferring the month to which the report correspond to
        from report data.
        "date" should be in form "yyyy-mm-dd"

        @return [str(year), str(month)]
        '''
        #assuming that all required contents had been filled
        date = report[list(report.keys())[0]][0]["date"].split("-")
        return date[0:2]

    def __add_header(self):
        '''
        use a paragraph to add a header. So call this before add
        other contents
        '''
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(GEN_REPORT.header_format % (self.month[0], self.month[1]))
        run.font.size = Pt(26)

    def __add_body(self, report):
        self.__add_intro()

        for item in report.keys():
            head = self.doc.add_paragraph()
            head.add_run(FEE_NAMES[item])
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_rows = len(report.keys()) + 1
            if(item == "SALARY_FEE"):
                num_cols = 4
                tab = self.doc.add_table(rows = num_rows, cols = num_cols)
                tab.style.font.size = self.__class__.table_font_size
                
                cell = tab.cell(0,0)
                cell.text = "日期"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell = tab.cell(0,1)
                cell.text = "姓名"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell = tab.cell(0,2)
                cell.text = "职称"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell = tab.cell(0,3)
                cell.text = "金额"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                num_cols = 3
                tab = self.doc.add_table(rows = num_rows, cols = num_cols)
                tab.style.font.size = self.__class__.table_font_size
                
                cell = tab.cell(0,0)
                cell.text = "日期"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell = tab.cell(0,1)
                cell.text = "项目名称"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell = tab.cell(0,2)
                cell.text = "金额"
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for idx in range(len(report[item])):
                if(item == "SALARY_FEE"):
                    cell = tab.cell(idx + 1,0)
                    cell.text = report[item][idx]["date"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell = tab.cell(idx + 1,1)
                    cell.text = report[item][idx]["name"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell = tab.cell(idx + 1,2)
                    cell.text = report[item][idx]["title"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell = tab.cell(idx + 1,3)
                    cell.text = report[item][idx]["amount"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell = tab.cell(idx + 1,0)
                    cell.text = report[item][idx]["date"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell = tab.cell(idx + 1,1)
                    cell.text = report[item][idx]["name"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell = tab.cell(idx + 1,2)
                    cell.text = report[item][idx]["amount"]
                    cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def __add_intro(self):
        p = self.doc.add_paragraph()
        run = p.add_run(GEN_REPORT.introduction)
        run.font.size = Pt(self.__class__.content_font_size)

if __name__ == '__main__':
    import test_suite
    report = GEN_REPORT(test_suite.REPORT_EXAMPLE)
    report.save("a.docx")

